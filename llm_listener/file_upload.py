"""File upload API with text extraction for document context."""

import hashlib
import io
import time
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any
from fastapi import APIRouter, HTTPException, UploadFile, File
from pydantic import BaseModel

router = APIRouter(prefix="/api/files", tags=["files"])

# In-memory file store with expiration
_file_store: Dict[str, dict] = {}
EXPIRATION_HOURS = 1
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

ALLOWED_TYPES = {
    'application/pdf': 'pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'text/plain': 'txt',
    'text/markdown': 'md',
    'text/xml': 'xml',
    'application/xml': 'xml',
    'application/json': 'json',
}


class FileUploadResponse(BaseModel):
    file_id: str
    filename: str
    content_type: str
    size: int
    extracted_chars: int
    preview: str


class FileContextResponse(BaseModel):
    file_id: str
    filename: str
    text: str


class FilesUploadResponse(BaseModel):
    files: List[FileUploadResponse]


def cleanup_expired():
    """Remove expired files from store."""
    now = time.time()
    expired = [fid for fid, data in _file_store.items()
               if data.get('expires_at', 0) < now]
    for fid in expired:
        del _file_store[fid]


def extract_pdf(content: bytes) -> str:
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return '\n\n'.join(text_parts)
    except ImportError:
        return "[PDF extraction requires pypdf library]"
    except Exception as e:
        return f"[Error extracting PDF: {str(e)}]"


def extract_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return '\n\n'.join(text_parts)
    except ImportError:
        return "[DOCX extraction requires python-docx library]"
    except Exception as e:
        return f"[Error extracting DOCX: {str(e)}]"


def extract_xml(content: bytes) -> str:
    """Extract text content from XML."""
    try:
        from xml.etree import ElementTree as ET

        def get_text(elem):
            texts = []
            if elem.text and elem.text.strip():
                texts.append(elem.text.strip())
            for child in elem:
                texts.extend(get_text(child))
            if elem.tail and elem.tail.strip():
                texts.append(elem.tail.strip())
            return texts

        root = ET.fromstring(content)
        texts = get_text(root)
        return '\n'.join(texts)
    except Exception as e:
        return f"[Error extracting XML: {str(e)}]"


def extract_json(content: bytes) -> str:
    """Pretty-print JSON content."""
    try:
        import json
        data = json.loads(content.decode('utf-8'))
        return json.dumps(data, indent=2)
    except Exception as e:
        return f"[Error extracting JSON: {str(e)}]"


def extract_text(content: bytes) -> str:
    """Extract plain text with encoding fallback."""
    try:
        return content.decode('utf-8')
    except UnicodeDecodeError:
        try:
            return content.decode('latin-1')
        except Exception:
            return content.decode('utf-8', errors='ignore')


def extract_content(content: bytes, content_type: str) -> str:
    """Extract text based on content type."""
    file_type = ALLOWED_TYPES.get(content_type, 'txt')

    if file_type == 'pdf':
        return extract_pdf(content)
    elif file_type == 'docx':
        return extract_docx(content)
    elif file_type == 'xml':
        return extract_xml(content)
    elif file_type == 'json':
        return extract_json(content)
    else:
        return extract_text(content)


@router.post("/upload", response_model=FilesUploadResponse)
async def upload_files(files: List[UploadFile] = File(...)):
    """Upload one or more files and extract text content."""
    cleanup_expired()

    results = []
    for file in files:
        # Validate content type
        if file.content_type not in ALLOWED_TYPES:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file.content_type}. Supported: PDF, DOCX, TXT, XML, JSON"
            )

        # Read content
        content = await file.read()

        # Validate size
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"File too large: {file.filename}. Max size: {MAX_FILE_SIZE // 1024 // 1024}MB"
            )

        # Generate file ID
        file_id = hashlib.sha256(content + file.filename.encode()).hexdigest()[:16]

        # Extract text
        extracted_text = extract_content(content, file.content_type)

        # Store with expiration
        expires_at = time.time() + (EXPIRATION_HOURS * 3600)
        _file_store[file_id] = {
            'filename': file.filename,
            'content_type': file.content_type,
            'size': len(content),
            'text': extracted_text,
            'expires_at': expires_at,
            'created_at': datetime.utcnow().isoformat(),
        }

        # Create preview (first 500 chars)
        preview = extracted_text[:500]
        if len(extracted_text) > 500:
            preview += '...'

        results.append(FileUploadResponse(
            file_id=file_id,
            filename=file.filename,
            content_type=file.content_type,
            size=len(content),
            extracted_chars=len(extracted_text),
            preview=preview,
        ))

    return FilesUploadResponse(files=results)


@router.get("/{file_id}/context", response_model=FileContextResponse)
async def get_file_context(file_id: str):
    """Get the extracted text content for a file."""
    cleanup_expired()

    if file_id not in _file_store:
        raise HTTPException(status_code=404, detail="File not found or expired")

    data = _file_store[file_id]
    return FileContextResponse(
        file_id=file_id,
        filename=data['filename'],
        text=data['text'],
    )


@router.delete("/{file_id}")
async def delete_file(file_id: str):
    """Delete an uploaded file."""
    if file_id in _file_store:
        del _file_store[file_id]
        return {"status": "deleted"}
    raise HTTPException(status_code=404, detail="File not found")


@router.get("/list")
async def list_files():
    """List all active (non-expired) files."""
    cleanup_expired()

    return {
        "files": [
            {
                "file_id": fid,
                "filename": data['filename'],
                "content_type": data['content_type'],
                "size": data['size'],
                "extracted_chars": len(data['text']),
                "created_at": data['created_at'],
            }
            for fid, data in _file_store.items()
        ]
    }


def build_file_context_prompt(question: str, file_contexts: List[Dict[str, Any]], max_chars: int = 50000) -> str:
    """Build prompt with file context for LLM query.

    Args:
        question: The user's question
        file_contexts: List of {filename, text} dicts
        max_chars: Maximum characters per document

    Returns:
        Enhanced prompt with document context
    """
    if not file_contexts:
        return question

    context_parts = ["You have been provided with the following document(s) as context:\n"]
    total_chars = 0

    for ctx in file_contexts:
        filename = ctx.get('filename', 'Document')
        text = ctx.get('text', '')

        # Truncate if needed
        if len(text) > max_chars:
            text = text[:max_chars] + f"\n\n[...truncated at {max_chars} characters...]"

        context_parts.append(f"\n{'='*50}")
        context_parts.append(f"Document: {filename}")
        context_parts.append('='*50)
        context_parts.append(text)
        total_chars += len(text)

    context_parts.append(f"\n{'='*50}\n")
    context_parts.append("Based on the above document(s) and your knowledge, please address:\n")
    context_parts.append(question)
    context_parts.append("\n\nIf the documents are relevant, cite specific information from them.")

    return '\n'.join(context_parts)


def get_file_context_by_id(file_id: str) -> Optional[Dict[str, str]]:
    """Get file context by ID for query injection."""
    cleanup_expired()
    if file_id in _file_store:
        data = _file_store[file_id]
        return {
            'filename': data['filename'],
            'text': data['text'],
        }
    return None
